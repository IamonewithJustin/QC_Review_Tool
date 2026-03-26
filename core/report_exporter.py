"""Export QC analysis results to a Word (.docx) document."""

from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Dict


def export_report(
    results: Dict[str, str],
    document_filename: str,
    output_path: str,
) -> None:
    """
    Write a Word document containing the QC analysis results.

    Args:
        results: mapping of model_name -> analysis_text
        document_filename: original reviewed file name (used in title)
        output_path: full path where the .docx should be saved
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        raise ImportError("python-docx is required. Run: pip install python-docx")

    doc = Document()

    # ── Patch every named style that carries spacing ────────────────────
    _zero_style_spacing(doc)

    # ── Helper: force single spacing + zero before/after on any paragraph
    def _ss(p) -> None:
        pf = p.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)

    # ── Title block ─────────────────────────────────────────────────────
    date_str = datetime.now().strftime("%B %d, %Y")

    title = doc.add_heading("QC Review Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _ss(title)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _ss(subtitle)
    run = subtitle.add_run(f"Document: {document_filename}    |    Date: {date_str}")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    gap = doc.add_paragraph()
    _ss(gap)

    # ── Body ────────────────────────────────────────────────────────────
    if not results:
        p = doc.add_paragraph("No analysis results available.")
        _ss(p)
    else:
        for model_name, analysis_text in results.items():
            h = doc.add_heading(f"Analysis by: {model_name}", level=1)
            _ss(h)
            _add_body_text(doc, analysis_text, _ss)
            spacer = doc.add_paragraph()
            _ss(spacer)

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


# ── Helpers ─────────────────────────────────────────────────────────────

def _zero_style_spacing(doc) -> None:
    """
    Remove space-before/after and enforce single-line spacing on every
    paragraph style in the document, including the document default (docDefaults).
    This prevents Word from re-applying style-level spacing even when paragraph
    direct-formatting says zero.
    """
    from docx.shared import Pt
    from docx.enum.text import WD_LINE_SPACING

    # Patch document-level default paragraph properties (docDefaults → pPrDefault)
    try:
        settings = doc.settings.element
        docDefaults = settings.find(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}docDefaults"
        )
        if docDefaults is not None:
            pPrDefault = docDefaults.find(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPrDefault"
            )
            if pPrDefault is not None:
                pPr = pPrDefault.find(
                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr"
                )
                if pPr is not None:
                    _set_spacing_xml(pPr)
    except Exception:
        pass

    # Patch every named paragraph style
    for style in doc.styles:
        try:
            if style.type.name != "PARAGRAPH":
                continue
            pf = style.paragraph_format
            pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
        except Exception:
            pass


def _set_spacing_xml(pPr_element) -> None:
    """
    Write a <w:spacing> element with single-line spacing, 0 before, 0 after
    directly into a pPr XML node (belt-and-suspenders for docDefaults).
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    existing = pPr_element.find(f"{{{W}}}spacing")
    if existing is not None:
        pPr_element.remove(existing)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    spacing.set(qn("w:line"), "240")       # 240 twips = 1 line (single)
    spacing.set(qn("w:lineRule"), "auto")
    pPr_element.append(spacing)


def _add_body_text(doc, text: str, single_space) -> None:
    """
    Add body text preserving paragraph breaks.
    Consecutive blank lines are collapsed to one blank paragraph to avoid
    visual double-spacing from stacked empty paragraphs.
    """
    lines = text.split("\n")
    prev_blank = False
    for line in lines:
        stripped = line.rstrip()
        if stripped:
            p = doc.add_paragraph(stripped)
            p.style = doc.styles["Normal"]
            single_space(p)
            prev_blank = False
        else:
            if not prev_blank:
                p = doc.add_paragraph()
                single_space(p)
            prev_blank = True
